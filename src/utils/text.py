import pdfminer
import re
import os
import io
import logging
import shutil
import subprocess
from typing import Optional, Dict, Any, List, Tuple

from pdfminer.high_level import extract_text, extract_pages
from pdfminer.layout import LTTextContainer


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    OCR_AVAILABLE = True
    logger.info("Image OCR libraries loaded successfully")
except ImportError as e:
    OCR_AVAILABLE = False
    logger.warning(f"Image OCR libraries not available: {e}. Image-based parsing may be limited.")

try:
    from pdf2image import convert_from_path
    PDF_IMAGE_CONVERSION_AVAILABLE = True
except ImportError as e:
    PDF_IMAGE_CONVERSION_AVAILABLE = False
    convert_from_path = None
    logger.warning(f"pdf2image not available: {e}. OCR fallback for PDF files will be disabled.")

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError as e:
    fitz = None
    PYMUPDF_AVAILABLE = False
    logger.warning(f"PyMuPDF not available: {e}. Hybrid per-page OCR routing will be disabled.")


def _resolve_tesseract_command() -> Optional[str]:
    env_cmd = os.getenv("TESSERACT_CMD")
    if env_cmd and os.path.isfile(env_cmd):
        return env_cmd

    path_cmd = shutil.which("tesseract")
    if path_cmd:
        return path_cmd

    common_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for cmd in common_paths:
        if os.path.isfile(cmd):
            return cmd
    return None


def _configure_tesseract() -> None:
    if not OCR_AVAILABLE:
        return

    resolved = _resolve_tesseract_command()
    if resolved:
        pytesseract.pytesseract.tesseract_cmd = resolved


_configure_tesseract()


try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    logger.warning(f"python-docx not available: {e}. DOCX parsing will be unavailable.")


def preprocess_image_for_ocr(image):
    try:
        
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        
        image = image.convert('L')
        
        
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        
        
        image = image.filter(ImageFilter.SHARPEN)
        
        
        image = image.point(lambda x: 0 if x < 128 else 255, '1')
        image = image.convert('L')
        
        return image
    except Exception as e:
        logger.warning(f"Image preprocessing failed: {e}")
        return image


def _median(values: List[float]) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    mid = len(ordered) // 2
    if len(ordered) % 2 == 0:
        return (ordered[mid - 1] + ordered[mid]) / 2.0
    return ordered[mid]


def _score_ocr_candidate(text: str) -> float:
    if not text or not text.strip():
        return 0.0

    base = calculate_text_confidence(text)

    heading_hits = 0
    section_terms = (
        "experience",
        "education",
        "skills",
        "summary",
        "projects",
        "certifications",
        "languages",
        "contact",
        "profile",
    )
    lowered = text.lower()
    for term in section_terms:
        if re.search(rf"\b{re.escape(term)}\b", lowered):
            heading_hits += 1

    heading_bonus = min(heading_hits * 0.03, 0.2)

    noisy_runs = len(re.findall(r"[^\w\s]{4,}", text))
    noise_penalty = min(noisy_runs * 0.02, 0.15)

    line_count = len([ln for ln in text.splitlines() if ln.strip()])
    structure_bonus = min(line_count / 100.0, 0.08)

    return round(max(base + heading_bonus + structure_bonus - noise_penalty, 0.0), 4)


def _group_lines_from_ocr_data(data: Dict[str, List[Any]]) -> List[Dict[str, Any]]:
    if not data or "text" not in data:
        return []

    grouped: Dict[tuple, Dict[str, Any]] = {}
    n = len(data.get("text", []))

    for i in range(n):
        word = str(data["text"][i]).strip()
        if not word:
            continue

        conf_raw = data.get("conf", [])[i] if i < len(data.get("conf", [])) else "-1"
        try:
            conf = float(conf_raw)
        except (TypeError, ValueError):
            continue
        if conf < 0:
            continue

        left = int(data.get("left", [0])[i]) if i < len(data.get("left", [])) else 0
        top = int(data.get("top", [0])[i]) if i < len(data.get("top", [])) else 0
        width = int(data.get("width", [0])[i]) if i < len(data.get("width", [])) else 0
        height = int(data.get("height", [0])[i]) if i < len(data.get("height", [])) else 0

        block_num = int(data.get("block_num", [0])[i]) if i < len(data.get("block_num", [])) else 0
        par_num = int(data.get("par_num", [0])[i]) if i < len(data.get("par_num", [])) else 0
        line_num = int(data.get("line_num", [0])[i]) if i < len(data.get("line_num", [])) else 0
        key = (block_num, par_num, line_num)

        if key not in grouped:
            grouped[key] = {
                "words": [],
                "left": left,
                "right": left + width,
                "top": top,
                "bottom": top + height,
            }

        grouped[key]["words"].append((left, word))
        grouped[key]["left"] = min(grouped[key]["left"], left)
        grouped[key]["right"] = max(grouped[key]["right"], left + width)
        grouped[key]["top"] = min(grouped[key]["top"], top)
        grouped[key]["bottom"] = max(grouped[key]["bottom"], top + height)

    lines: List[Dict[str, Any]] = []
    for item in grouped.values():
        words = [w for _, w in sorted(item["words"], key=lambda x: x[0])]
        text = " ".join(words).strip()
        if not text:
            continue
        lines.append({
            "text": text,
            "left": item["left"],
            "right": item["right"],
            "top": item["top"],
            "bottom": item["bottom"],
            "center_x": (item["left"] + item["right"]) / 2.0,
        })

    return lines


def _reconstruct_column_aware_text(lines: List[Dict[str, Any]], page_width: int) -> str:
    if not lines:
        return ""

    lines = sorted(lines, key=lambda l: (l["top"], l["left"]))
    left_candidates = [ln for ln in lines if ln["center_x"] < page_width * 0.46]
    right_candidates = [ln for ln in lines if ln["center_x"] > page_width * 0.54]

    is_two_column = len(left_candidates) >= 5 and len(right_candidates) >= 5
    if not is_two_column:
        return "\n".join(ln["text"] for ln in lines)

    left_centers = [ln["center_x"] for ln in left_candidates]
    right_centers = [ln["center_x"] for ln in right_candidates]
    separation = _median(right_centers) - _median(left_centers)
    if separation < page_width * 0.2:
        return "\n".join(ln["text"] for ln in lines)

    split_x = (_median(left_centers) + _median(right_centers)) / 2.0
    line_heights = [max(1, ln["bottom"] - ln["top"]) for ln in lines]
    avg_line_height = _median([float(h) for h in line_heights])

    left_col: List[Dict[str, Any]] = []
    right_col: List[Dict[str, Any]] = []
    for ln in lines:
        if ln["center_x"] <= split_x:
            left_col.append(ln)
        else:
            right_col.append(ln)

    if not left_col or not right_col:
        return "\n".join(ln["text"] for ln in lines)

    column_start_top = min(
        min(ln["top"] for ln in left_col),
        min(ln["top"] for ln in right_col),
    )

    preamble: List[Dict[str, Any]] = []
    left_body: List[Dict[str, Any]] = []
    right_body: List[Dict[str, Any]] = []

    for ln in lines:
        if ln["top"] < column_start_top - avg_line_height:
            preamble.append(ln)
            continue

        if ln["center_x"] <= split_x:
            left_body.append(ln)
        else:
            right_body.append(ln)

    preamble = sorted(preamble, key=lambda l: (l["top"], l["left"]))
    left_body = sorted(left_body, key=lambda l: (l["top"], l["left"]))
    right_body = sorted(right_body, key=lambda l: (l["top"], l["left"]))

    ordered = [ln["text"] for ln in preamble] + [ln["text"] for ln in left_body] + [ln["text"] for ln in right_body]
    return "\n".join(ordered)


def _extract_image_text_with_layout_aware_ocr(image) -> str:
    candidates: List[str] = []
    for psm in (1, 3, 4, 6):
        config = f"--oem 3 --psm {psm}"
        try:
            candidate = pytesseract.image_to_string(image, config=config)
            if candidate and candidate.strip():
                candidates.append(candidate)
        except Exception:
            continue

    try:
        data = pytesseract.image_to_data(
            image,
            config="--oem 3 --psm 1",
            output_type=pytesseract.Output.DICT,
        )
        lines = _group_lines_from_ocr_data(data)
        reconstructed = _reconstruct_column_aware_text(lines, image.width)
        if reconstructed and reconstructed.strip():
            candidates.append(reconstructed)
    except Exception as e:
        logger.debug(f"Could not perform column-aware OCR reconstruction: {e}")

    if not candidates:
        return ""

    best_text = ""
    best_score = -1.0
    for text in candidates:
        score = _score_ocr_candidate(text)
        if score > best_score:
            best_score = score
            best_text = text

    return best_text


def extract_text_with_ocr(pdf_path, use_preprocessing=True):
    if not OCR_AVAILABLE or not PDF_IMAGE_CONVERSION_AVAILABLE:
        logger.warning("PDF OCR unavailable (missing pytesseract/Pillow or pdf2image)")
        return ""

    try:
        logger.info(f"Starting OCR extraction for: {pdf_path}")
        
        
        images = convert_from_path(pdf_path, dpi=300)
        
        text = ""
        for page_num, image in enumerate(images):
            logger.info(f"Processing page {page_num + 1} with OCR")
            
            
            if use_preprocessing:
                image = preprocess_image_for_ocr(image)
            
            
            custom_config = r'--oem 3 --psm 6'  
            page_text = pytesseract.image_to_string(image, config=custom_config)
            text += page_text + "\n"
        
        logger.info(f"OCR extraction completed. Text length: {len(text)}")
        return text
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return ""


def extract_text_with_ocr_fallback(pdf_path):
    if not OCR_AVAILABLE or not PDF_IMAGE_CONVERSION_AVAILABLE:
        return ""

    
    text = extract_text_with_ocr(pdf_path, use_preprocessing=True)
    if text and len(text.strip()) > 50:
        return text
    
    
    logger.info("Retrying OCR without preprocessing")
    text = extract_text_with_ocr(pdf_path, use_preprocessing=False)
    if text and len(text.strip()) > 50:
        return text
    
    
    logger.info("Retrying OCR with different configuration")
    try:
        images = convert_from_path(pdf_path, dpi=300)
        text = ""
        for image in images:
            
            custom_config = r'--oem 3 --psm 1'
            page_text = pytesseract.image_to_string(image, config=custom_config)
            text += page_text + "\n"
        
        if text and len(text.strip()) > 50:
            return text
    except Exception as e:
        logger.error(f"Fallback OCR also failed: {e}")
    
    return ""


def calculate_text_confidence(text):
    if not text or not text.strip():
        return 0.0
    
    text_length = len(text.strip())
    word_count = len(text.split())
    
    
    error_patterns = [
        r'[\x00-\x08\x0b-\x0c\x0e-\x1f]',  
        r'[_\~]{5,}',  
    ]
    
    error_count = 0
    for pattern in error_patterns:
        error_count += len(re.findall(pattern, text))
    
    
    
    length_factor = min(text_length / 1000, 1.0) * 0.4
    
    
    word_factor = min(word_count / 200, 1.0) * 0.3
    
    
    error_factor = max(1.0 - (error_count / 100), 0.0) * 0.3
    
    confidence = length_factor + word_factor + error_factor
    
    return round(confidence, 2)


def _extract_pdfminer_text_by_page(pdf_path: str) -> List[str]:
    pages: List[str] = []
    for page in extract_pages(pdf_path):
        chunks: List[str] = []
        for element in page:
            if isinstance(element, LTTextContainer):
                text = element.get_text().strip()
                if text:
                    chunks.append(text)
        pages.append("\n".join(chunks).strip())
    return pages


def _extract_pdfminer_column_aware_text(pdf_path: str) -> Tuple[str, float, bool]:
    pages: List[str] = []
    two_column_detected = False

    for page in extract_pages(pdf_path):
        page_width = int(getattr(page, "width", 612))
        page_height = float(getattr(page, "height", 792))
        layout_lines: List[Dict[str, Any]] = []

        for element in page:
            if not isinstance(element, LTTextContainer):
                continue
            block_text = (element.get_text() or "").strip()
            if not block_text:
                continue
            x0, y0, x1, y1 = element.bbox
            center_x = (x0 + x1) / 2.0
            top = max(0.0, page_height - y1)
            bottom = max(top + 1.0, page_height - y0)

            block_lines = [ln.strip() for ln in block_text.splitlines() if ln.strip()]
            for line in block_lines:
                layout_lines.append({
                    "text": line,
                    "left": int(x0),
                    "right": int(x1),
                    "top": int(top),
                    "bottom": int(bottom),
                    "center_x": float(center_x),
                })

        if not layout_lines:
            pages.append("")
            continue

        left_candidates = [ln for ln in layout_lines if ln["center_x"] < page_width * 0.46]
        right_candidates = [ln for ln in layout_lines if ln["center_x"] > page_width * 0.54]
        if len(left_candidates) >= 5 and len(right_candidates) >= 5:
            two_column_detected = True

        page_text = _reconstruct_column_aware_text(layout_lines, page_width).strip()
        pages.append(page_text)

    combined = "\n\n".join([p for p in pages if p]).strip()
    if not combined:
        return "", 0.0, two_column_detected
    return combined, calculate_text_confidence(combined), two_column_detected


def _extract_text_with_page_routed_ocr(pdf_path: str) -> Tuple[str, float]:
    if not (OCR_AVAILABLE and PYMUPDF_AVAILABLE):
        return "", 0.0

    try:
        page_texts = _extract_pdfminer_text_by_page(pdf_path)
        if not page_texts:
            return "", 0.0

        doc = fitz.open(pdf_path)
        merged_pages: List[str] = []
        page_confidences: List[float] = []
        ocr_pages = 0
        pdfminer_pages = 0

        try:
            for idx, page_text in enumerate(page_texts):
                cleaned_page_text = (page_text or "").strip()
                miner_conf = calculate_text_confidence(cleaned_page_text) if cleaned_page_text else 0.0

                # Route each page independently: keep strong text pages, OCR only weak/empty pages.
                should_try_ocr = (not cleaned_page_text) or (len(cleaned_page_text) < 80) or (miner_conf < 0.28)
                best_text = cleaned_page_text
                best_conf = miner_conf

                if should_try_ocr and idx < len(doc):
                    try:
                        page = doc[idx]
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        image = Image.open(io.BytesIO(pix.tobytes("png")))
                        image = preprocess_image_for_ocr(image)
                        ocr_text = _extract_image_text_with_layout_aware_ocr(image)
                        ocr_conf = calculate_text_confidence(ocr_text) * 0.85 if ocr_text else 0.0

                        if ocr_text and (ocr_conf >= best_conf + 0.08 or len(ocr_text.strip()) > len(best_text.strip()) * 1.2):
                            best_text = ocr_text.strip()
                            best_conf = min(ocr_conf, 1.0)
                            ocr_pages += 1
                        else:
                            pdfminer_pages += 1
                    except Exception as page_ocr_error:
                        logger.debug(f"Per-page OCR failed for page {idx + 1}: {page_ocr_error}")
                        pdfminer_pages += 1
                else:
                    pdfminer_pages += 1

                if best_text:
                    merged_pages.append(best_text)
                    page_confidences.append(best_conf)
        finally:
            doc.close()

        if not merged_pages:
            return "", 0.0

        combined_text = "\n\n".join(merged_pages).strip()
        combined_confidence = round(sum(page_confidences) / len(page_confidences), 2) if page_confidences else 0.0
        logger.info(
            "Hybrid page-routed extraction completed. Pages: %s (pdfminer=%s, ocr=%s), text length=%s, confidence=%s",
            len(page_texts),
            pdfminer_pages,
            ocr_pages,
            len(combined_text),
            combined_confidence,
        )
        return combined_text, combined_confidence

    except Exception as e:
        logger.warning(f"Hybrid page-routed extraction failed: {e}")
        return "", 0.0


def extract_text_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        return "", 0.0

    ext = os.path.splitext(pdf_path)[1].lower()
    if ext != ".pdf":
        return extract_text_from_file(pdf_path)
    
    text = ""
    confidence = 0.0
    
    
    try:
        logger.info(f"Trying pdfminer extraction for: {pdf_path}")
        text = extract_text(pdf_path)

        column_text, column_confidence, two_column_detected = _extract_pdfminer_column_aware_text(pdf_path)
        if column_text and (column_confidence >= confidence + 0.04 or (two_column_detected and column_confidence >= 0.35)):
            logger.info(
                "Using column-aware PDFMiner reconstruction%s. Length: %s, Confidence: %s",
                " (two-column detected)" if two_column_detected else "",
                len(column_text),
                column_confidence,
            )
            text = column_text
            confidence = column_confidence

        if text and len(text.strip()) > 50:
            if confidence <= 0.0:
                confidence = calculate_text_confidence(text)
            logger.info(f"pdfminer extraction successful. Text length: {len(text)}, Confidence: {confidence}")
            if confidence >= 0.7 and len(text.strip()) >= 180:
                return text, confidence
            logger.info("pdfminer output looks partial/low-confidence; trying hybrid per-page routing")
        else:
            logger.warning(f"pdfminer returned empty or very short text. Length: {len(text) if text else 0}")
    except Exception as e:
        logger.error(f"pdfminer extraction failed: {e}")
    
    if OCR_AVAILABLE and PYMUPDF_AVAILABLE:
        hybrid_text, hybrid_confidence = _extract_text_with_page_routed_ocr(pdf_path)
        if hybrid_text and len(hybrid_text.strip()) > 50:
            if hybrid_confidence >= max(confidence, 0.3):
                logger.info(
                    "Using hybrid page-routed extraction. Length: %s, Confidence: %s",
                    len(hybrid_text),
                    hybrid_confidence,
                )
                return hybrid_text, hybrid_confidence
            logger.info(
                "Hybrid extraction produced lower confidence than current text. Keeping current best. Hybrid=%s Current=%s",
                hybrid_confidence,
                confidence,
            )
    
    
    if OCR_AVAILABLE and PDF_IMAGE_CONVERSION_AVAILABLE:
        logger.info("Falling back to OCR extraction")
        ocr_text = extract_text_with_ocr_fallback(pdf_path)
        
        if ocr_text and len(ocr_text.strip()) > 50:
            confidence = calculate_text_confidence(ocr_text) * 0.8  
            logger.info(f"OCR extraction successful. Text length: {len(ocr_text)}, Confidence: {confidence}")
            return ocr_text, confidence
        else:
            logger.warning("OCR also returned empty or short text")
    else:
        logger.warning("PDF OCR fallback not available. Returning pdfminer result even if empty.")
    
    
    confidence = calculate_text_confidence(text)
    return text if text else "", confidence


def extract_text_from_docx(docx_path):
    if not DOCX_AVAILABLE:
        logger.error("DOCX support unavailable because python-docx is not installed")
        return "", 0.0

    try:
        document = Document(docx_path)
        chunks = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in document.tables:
            for row in table.rows:
                row_values = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_values.append(cell_text)
                if row_values:
                    chunks.append(" | ".join(row_values))

        text = "\n".join(chunks).strip()
        confidence = calculate_text_confidence(text)
        if text:
            confidence = min(1.0, round(confidence + 0.1, 2))
        return text, confidence
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return "", 0.0


def _extract_text_from_doc_with_tool(doc_path: str, tool_name: str):
    if not shutil.which(tool_name):
        return ""
    try:
        proc = subprocess.run(
            [tool_name, doc_path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=30,
            check=False,
        )
        if proc.returncode == 0:
            return (proc.stdout or "").strip()
    except Exception as e:
        logger.debug(f"{tool_name} failed for DOC parsing: {e}")
    return ""


def extract_text_from_doc(doc_path):
    text = _extract_text_from_doc_with_tool(doc_path, "antiword")
    if not text:
        text = _extract_text_from_doc_with_tool(doc_path, "catdoc")

    confidence = calculate_text_confidence(text)
    if text:
        confidence = min(1.0, round(confidence * 0.9, 2))
    return text, confidence


def extract_text_from_file(file_path):
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return "", 0.0

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".doc":
        return extract_text_from_doc(file_path)
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return extract_text_from_image(file_path)

    logger.warning(f"Unsupported file type '{ext}' for extraction")
    return "", 0.0




def extract_text_from_image(image_path):
    if not OCR_AVAILABLE:
        logger.error("OCR libraries not available")
        return "", 0.0
    
    try:
        image = Image.open(image_path)
        
        
        image = preprocess_image_for_ocr(image)
        text = _extract_image_text_with_layout_aware_ocr(image)
        
        confidence = calculate_text_confidence(text)
        return text, confidence
    except Exception as e:
        logger.error(f"Image OCR extraction failed: {e}")
        return "", 0.0
 
 
if __name__ == '__main__':
    pdf_path = "sample_resume.pdf"
    text, confidence = extract_text_from_pdf(pdf_path)
    print(f"Text: {text[:500]}...")
    print(f"Confidence: {confidence}")
